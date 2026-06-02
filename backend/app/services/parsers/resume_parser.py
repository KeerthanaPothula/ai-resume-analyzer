import re
from typing import Optional
import fitz  # PyMuPDF
from docx import Document


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF."""
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text.strip()
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX."""
    try:
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def extract_email(text: str) -> Optional[str]:
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    matches = re.findall(pattern, text)
    return matches[0] if matches else None


def extract_phone(text: str) -> Optional[str]:
    pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    matches = re.findall(pattern, text)
    if matches:
        phone = ''.join(matches[0]) if isinstance(matches[0], tuple) else matches[0]
        return phone.strip()
    return None


def extract_name(text: str) -> Optional[str]:
    """Extract candidate name from first few lines."""
    lines = [ln.strip() for ln in text.split('\n') if ln.strip()]
    if lines:
        first_line = lines[0]
        # Likely a name if it's 2-4 words, no special chars, not all caps keywords
        words = first_line.split()
        if 2 <= len(words) <= 4 and all(w.replace('-', '').isalpha() for w in words):
            return first_line
    return None


def extract_location(text: str) -> Optional[str]:
    """Simple location extraction."""
    patterns = [
        r'\b([A-Z][a-z]+(?:\s[A-Z][a-z]+)?,\s*[A-Z]{2})\b',  # City, ST
        r'\b([A-Z][a-z]+(?:\s[A-Z][a-z]+)?,\s*[A-Z][a-z]+)\b',  # City, Country
    ]
    for pattern in patterns:
        matches = re.findall(pattern, text[:500])
        if matches:
            return matches[0]
    return None


def extract_experience_years(text: str) -> float:
    """Extract years of experience mentioned in resume."""
    patterns = [
        r'(\d+)\+?\s*years?\s+(?:of\s+)?experience',
        r'experience[:\s]+(\d+)\+?\s*years?',
        r'(\d+)\+?\s*yrs?\s+experience',
    ]
    for pattern in patterns:
        matches = re.findall(pattern, text.lower())
        if matches:
            return float(matches[0])

    # Estimate from date ranges
    date_pattern = r'\b(19|20)\d{2}\b'
    years = [int(y) for y in re.findall(date_pattern, text)]
    if len(years) >= 2:
        return float(max(years) - min(years))
    return 0.0


def extract_education_level(text: str) -> str:
    """Extract highest education level."""
    text_lower = text.lower()
    if any(kw in text_lower for kw in ['ph.d', 'phd', 'doctorate', 'doctoral']):
        return 'PhD'
    elif any(kw in text_lower for kw in ["master's", 'masters', 'mba', 'm.s.', 'msc', 'm.tech']):
        return "Master's"
    elif any(kw in text_lower for kw in ["bachelor's", 'bachelors', 'b.s.', 'b.e.', 'b.tech', 'bsc', 'undergraduate']):
        return "Bachelor's"
    elif any(kw in text_lower for kw in ['associate', 'diploma']):
        return 'Associate/Diploma'
    elif any(kw in text_lower for kw in ['high school', 'secondary', 'ged']):
        return 'High School'
    return 'Not Specified'


def parse_resume_file(file_path: str, file_type: str) -> dict:
    """Parse resume file and extract structured data."""
    if file_type == "pdf":
        text = extract_text_from_pdf(file_path)
    elif file_type in ["docx", "doc"]:
        text = extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

    return {
        "raw_text": text,
        "candidate_name": extract_name(text),
        "candidate_email": extract_email(text),
        "candidate_phone": extract_phone(text),
        "candidate_location": extract_location(text),
        "experience_years": extract_experience_years(text),
        "education_level": extract_education_level(text),
    }
